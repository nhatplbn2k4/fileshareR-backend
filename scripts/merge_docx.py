#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script để merge nhiều file docx theo thứ tự sử dụng docxcompose
Usage: python merge_docx.py <output.docx> <file1.docx:is_special> <file2.docx:is_special> ...
is_special: 0 = normal (pdf2docx), 1 = special (LibreOffice)
"""

import sys
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer

def add_section_break(doc):
    """Thêm section break vào cuối document với margins = 0"""
    # Tạo paragraph MỚI để chứa section break
    # Điều này ngăn section properties ảnh hưởng đến section trước đó
    last_para = doc.add_paragraph()
    
    # Add section properties to paragraph (creates section break)
    pPr = last_para._element.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    pPr.append(sectPr)
    
    # Set page margins to 0 để tránh nhảy trang
    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'), '0')
    pgMar.set(qn('w:right'), '0')
    pgMar.set(qn('w:bottom'), '0')
    pgMar.set(qn('w:left'), '0')
    pgMar.set(qn('w:header'), '0')
    pgMar.set(qn('w:footer'), '0')
    pgMar.set(qn('w:gutter'), '0')
    sectPr.append(pgMar)

def add_section_break_with_margins(doc, margins):
    """Thêm section break vào cuối document với margins được truyền vào"""
    # Tạo paragraph MỚI để chứa section break
    last_para = doc.add_paragraph()
    
    # Add section properties to paragraph (creates section break)
    pPr = last_para._element.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    pPr.append(sectPr)
    
    # Set margins từ dict được truyền vào
    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'), str(int(margins['top'])))
    pgMar.set(qn('w:right'), str(int(margins['right'])))
    pgMar.set(qn('w:bottom'), str(int(margins['bottom'])))
    pgMar.set(qn('w:left'), str(int(margins['left'])))
    pgMar.set(qn('w:header'), '720')
    pgMar.set(qn('w:footer'), '720')
    pgMar.set(qn('w:gutter'), '0')
    sectPr.append(pgMar)
    
def add_page_break(doc):
    """Thêm page break vào cuối document"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def merge_multiple_documents(output_path, file_infos):
    """
    Merge nhiều file docx vào 1 file (không normalize - để Java xử lý)
    file_infos: list of tuples (file_path, is_special)
    """
    if not file_infos:
        print("No files to merge")
        return False
    
    if len(file_infos) == 1:
        # Chỉ có 1 file, copy trực tiếp
        import shutil
        shutil.copy(file_infos[0][0], output_path)
        print(f"Single file copied: {output_path}")
        # Output paragraph count cho Java
        doc = Document(file_infos[0][0])
        print(f"File 0: {file_infos[0][0]} - sections={len(doc.sections)}, paragraphs={len(doc.paragraphs)}")
        print(f"PARA_COUNT:0:{len(doc.paragraphs)}")
        return True
    
    try:
        # Mở file đầu tiên làm base
        first_file, first_is_special = file_infos[0]
        master = Document(first_file)
        master_sections = len(master.sections)
        print(f"File 0: {first_file} - sections={master_sections}, paragraphs={len(master.paragraphs)}")
        
        composer = Composer(master)
        
        # Track paragraph counts từng file cho Java normalize
        para_counts = []
        para_counts.append(len(master.paragraphs))
        
        # Track 2 file segment gần nhất: [file trước, file hiện tại]
        # Ban đầu chỉ có master
        recent_files = [master]
        recent_files_info = [{'index': 0, 'is_special': first_is_special, 'path': first_file}]
        
        # Append các file còn lại
        for i in range(1, len(file_infos)):
            file_path, is_special = file_infos[i]
            
            # Load file hiện tại để lấy thông tin
            doc = Document(file_path)
            doc_sections = len(doc.sections)
            para_count = len(doc.paragraphs)
            print(f"File {i}: {file_path} - sections={doc_sections}, paragraphs={para_count}, special={is_special}")
            para_counts.append(para_count)
            
            # Thêm file hiện tại vào list
            recent_files.append(doc)
            recent_files_info.append({'index': i, 'is_special': is_special, 'path': file_path})
            
            # Giữ chỉ 2 file gần nhất
            if len(recent_files) > 2:
                recent_files.pop(0)
                recent_files_info.pop(0)
            
            # Lấy margins từ file TRƯỚC (file đầu tiên trong list)
            prev_file = recent_files[0]
            prev_info = recent_files_info[0]
            prev_section = prev_file.sections[-1]
            
            # Convert EMU (python-docx) sang twips (Word XML)
            # 1 inch = 914400 EMU = 1440 twips
            # twips = EMU / 635
            prev_margins = {
                'left': int(prev_section.left_margin / 635),
                'right': int(prev_section.right_margin / 635),
                'top': int(prev_section.top_margin / 635),
                'bottom': int(prev_section.bottom_margin / 635)
            }
            print(f"Previous file {prev_info['index']} (special={prev_info['is_special']}) margins: L={prev_margins['left']}, R={prev_margins['right']}, T={prev_margins['top']}, B={prev_margins['bottom']}")
            
            # Thêm page break
            add_page_break(composer.doc)
            
            # Thêm section break
            if not is_special:
                # File normal: margins = 0
                add_section_break(composer.doc)
                print(f"Added section break (margins=0) before file {i}")
            else:
                # File special: giữ margins từ file TRƯỚC
                add_section_break_with_margins(composer.doc, prev_margins)
                print(f"Added section break (keep prev margins) before special file {i}")
            
            composer.append(doc)
            print(f"Appended file {i}")




        
        # Save (không normalize - để Java xử lý)
        composer.save(output_path)
        print(f"Merged successfully: {output_path}")
        
        # Output paragraph counts cho Java parse
        for idx, count in enumerate(para_counts):
            print(f"PARA_COUNT:{idx}:{count}")
        
        return True
        
    except Exception as e:
        print(f"Error merging: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    if len(sys.argv) < 3:
        print("Usage: python merge_docx.py <output.docx> <file1.docx:is_special> <file2.docx:is_special> ...")
        print("is_special: 0 = normal (pdf2docx), 1 = special (LibreOffice)")
        sys.exit(1)
    
    output_path = sys.argv[1]
    file_infos = []
    
    for arg in sys.argv[2:]:
        if ':' in arg:
            parts = arg.rsplit(':', 1)
            file_path = parts[0]
            is_special = parts[1] == '1'
        else:
            file_path = arg
            is_special = False
        
        file_infos.append((file_path, is_special))
        print(f"File: {file_path}, special={is_special}")
    
    success = merge_multiple_documents(output_path, file_infos)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
